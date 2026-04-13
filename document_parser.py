"""Document parsing and writing utilities for DOCX and PDF files."""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import List

import pdfplumber
from docx import Document


@dataclass
class Block:
    """A lightweight structural block extracted from the source document."""

    kind: str
    text: str
    level: int = 1


REFERENCE_HEADING_RE = re.compile(r"^(参考文献|references|bibliography)\s*$", re.IGNORECASE)
REFERENCE_ENTRY_RE = re.compile(r"^(\[\d+\]|\(\d+\)|\d+\.)\s+")
CAPTION_RE = re.compile(
    r"^(图|表)\s*[0-9一二三四五六七八九十百]+|^(figure|table)\s*\d+[a-zA-Z]?",
    re.IGNORECASE,
)
SECTION_HEADING_RE = re.compile(r"^(\d+(\.\d+){0,3}|[一二三四五六七八九十]+[、.])\s*\S+")


def parse_uploaded_file(uploaded_file) -> List[Block]:
    """Parse an uploaded Streamlit file into structural blocks."""
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]
    file_bytes = uploaded_file.getvalue()

    if suffix == "docx":
        return parse_docx_bytes(file_bytes)
    if suffix == "pdf":
        return parse_pdf_bytes(file_bytes)

    raise ValueError("Unsupported file format. Please upload .docx or .pdf.")


def parse_docx_bytes(file_bytes: bytes) -> List[Block]:
    """Extract blocks from a DOCX binary."""
    doc = Document(BytesIO(file_bytes))
    blocks: List[Block] = []
    in_references = False

    for para in doc.paragraphs:
        raw_text = para.text or ""
        stripped = raw_text.strip()

        if not stripped:
            blocks.append(Block(kind="blank", text="", level=1))
            continue

        style_name = para.style.name if para.style else ""

        if _is_reference_heading(stripped):
            in_references = True
            blocks.append(Block(kind="heading", text=stripped, level=2))
            continue

        if in_references:
            if style_name.lower().startswith("heading") and not _is_reference_heading(stripped):
                in_references = False
            else:
                blocks.append(Block(kind="reference", text=stripped, level=1))
                continue

        if style_name.lower().startswith("heading"):
            blocks.append(
                Block(kind="heading", text=stripped, level=_extract_heading_level(style_name))
            )
        elif style_name.lower() == "caption" or _is_caption(stripped):
            blocks.append(Block(kind="caption", text=stripped, level=1))
        elif _looks_like_heading(stripped):
            blocks.append(Block(kind="heading", text=stripped, level=2))
        else:
            blocks.append(Block(kind="paragraph", text=stripped, level=1))

    return blocks


def parse_pdf_bytes(file_bytes: bytes) -> List[Block]:
    """Extract blocks from a PDF binary using plain-text heuristics."""
    blocks: List[Block] = []
    in_references = False

    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if not text.strip():
                continue

            paragraphs = _split_pdf_text_to_paragraphs(text)
            for para in paragraphs:
                stripped = para.strip()
                if not stripped:
                    blocks.append(Block(kind="blank", text="", level=1))
                    continue

                if _is_reference_heading(stripped):
                    in_references = True
                    blocks.append(Block(kind="heading", text=stripped, level=2))
                    continue

                if in_references:
                    if _looks_like_heading(stripped) and not _looks_like_reference_entry(stripped):
                        in_references = False
                    else:
                        blocks.append(Block(kind="reference", text=stripped, level=1))
                        continue

                if _is_caption(stripped):
                    blocks.append(Block(kind="caption", text=stripped, level=1))
                elif _looks_like_heading(stripped):
                    blocks.append(Block(kind="heading", text=stripped, level=2))
                elif _looks_like_reference_entry(stripped):
                    blocks.append(Block(kind="reference", text=stripped, level=1))
                else:
                    blocks.append(Block(kind="paragraph", text=stripped, level=1))

            blocks.append(Block(kind="blank", text="", level=1))

    return _trim_edge_blanks(blocks)


def blocks_to_docx_bytes(blocks: List[Block]) -> bytes:
    """Write blocks into a DOCX binary."""
    doc = Document()

    for block in blocks:
        if block.kind == "blank":
            doc.add_paragraph("")
            continue

        if block.kind == "heading":
            level = min(max(block.level, 1), 9)
            doc.add_heading(block.text, level=level)
            continue

        paragraph = doc.add_paragraph(block.text)
        if block.kind == "caption":
            try:
                paragraph.style = "Caption"
            except KeyError:
                pass

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def count_characters(blocks: List[Block]) -> int:
    """Count characters in extracted blocks."""
    return sum(len(block.text) for block in blocks)


def _extract_heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)", style_name)
    if not match:
        return 1
    return min(max(int(match.group(1)), 1), 9)


def _is_reference_heading(text: str) -> bool:
    return bool(REFERENCE_HEADING_RE.match(text.strip()))


def _looks_like_reference_entry(text: str) -> bool:
    return bool(REFERENCE_ENTRY_RE.match(text.strip()))


def _is_caption(text: str) -> bool:
    return bool(CAPTION_RE.match(text.strip()))


def _looks_like_heading(text: str) -> bool:
    stripped = text.strip()
    if len(stripped) > 120:
        return False
    if SECTION_HEADING_RE.match(stripped):
        return True
    if stripped.isupper() and 4 <= len(stripped) <= 80:
        return True
    if stripped.endswith((":", "：")) and len(stripped) <= 80:
        return True
    return False


def _split_pdf_text_to_paragraphs(text: str) -> List[str]:
    lines = text.splitlines()
    paragraphs: List[str] = []
    buffer: List[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if buffer:
                paragraphs.append(_merge_wrapped_lines(buffer))
                buffer = []
            paragraphs.append("")
            continue
        buffer.append(line)

    if buffer:
        paragraphs.append(_merge_wrapped_lines(buffer))

    return paragraphs


def _merge_wrapped_lines(lines: List[str]) -> str:
    merged = lines[0]
    for nxt in lines[1:]:
        if _should_concat_without_space(merged, nxt):
            merged += nxt
        else:
            merged += f" {nxt}"
    return merged


def _should_concat_without_space(left: str, right: str) -> bool:
    if not left or not right:
        return False

    left_last = left[-1]
    right_first = right[0]

    if _is_cjk(left_last) and _is_cjk(right_first):
        return True

    if left_last in "-‐‑‒–—" and right_first.isalnum():
        return True

    return False


def _is_cjk(char: str) -> bool:
    code = ord(char)
    return 0x4E00 <= code <= 0x9FFF


def _trim_edge_blanks(blocks: List[Block]) -> List[Block]:
    start = 0
    end = len(blocks)

    while start < end and blocks[start].kind == "blank":
        start += 1

    while end > start and blocks[end - 1].kind == "blank":
        end -= 1

    return blocks[start:end]
